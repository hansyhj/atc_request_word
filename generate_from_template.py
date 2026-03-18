#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import logging
import re
import tempfile
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable

import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)


TEMPLATE_MAIN_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
DOC_MAIN_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"


@dataclass(frozen=True)
class Proto:
    p_style: str
    ppr_xml: object | None
    rpr_xml: object | None


def _validate_inputs(template: Path, content: Path) -> None:
    if not template.exists():
        raise FileNotFoundError(f"模板文件不存在：{template}")
    if template.suffix.lower() != ".dotx":
        raise ValueError(f"模板文件必须是 .dotx 格式，当前：{template.suffix}")
    if not content.exists():
        raise FileNotFoundError(f"内容 JSON 不存在：{content}")


def _load_json(path: Path) -> dict:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        raise ValueError(f"JSON 格式错误（{path}）：{e}") from e
    if not isinstance(data, dict):
        raise ValueError(f"JSON 顶层必须是对象，当前类型：{type(data).__name__}")
    return data


def _dotx_to_docx(dotx: Path, tmpdir: Path) -> Path:
    """
    python-docx does not open .dotx directly (template content-type).
    Convert by patching [Content_Types].xml and writing a temp .docx.
    tmpdir is managed by the caller (use tempfile.TemporaryDirectory).
    """
    out = tmpdir / (dotx.stem + ".docx")
    with zipfile.ZipFile(dotx, "r") as zin, zipfile.ZipFile(
        out, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == "[Content_Types].xml":
                root = ET.fromstring(data)
                ns = {"ct": "http://schemas.openxmlformats.org/package/2006/content-types"}
                for ov in root.findall("ct:Override", ns):
                    if ov.attrib.get("PartName") == "/word/document.xml":
                        ct = ov.attrib.get("ContentType", "")
                        if ct == TEMPLATE_MAIN_CT:
                            ov.set("ContentType", DOC_MAIN_CT)
                data = ET.tostring(
                    root, encoding="utf-8", xml_declaration=True, short_empty_elements=True
                )
            zout.writestr(info, data)
    return out


def _iter_nonempty_paras(doc: Document) -> Iterable:
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            yield p


def _find_first(doc: Document, predicate) -> object:
    for p in _iter_nonempty_paras(doc):
        if predicate(p):
            return p
    raise ValueError("Cannot find required paragraph prototype in template.")


def _find_index(doc: Document, paragraph) -> int:
    target = paragraph._p  # pylint: disable=protected-access
    for i, p in enumerate(doc.paragraphs):
        if p._p is target:  # pylint: disable=protected-access
            return i
    return -1


def _proto_from_para(p) -> Proto:
    ppr = p._p.pPr  # pylint: disable=protected-access
    rpr = None
    for r in p.runs:
        if r.text is None:
            continue
        rpr = r._r.rPr  # pylint: disable=protected-access
        if rpr is not None:
            break
    return Proto(
        p_style=p.style.name if p.style is not None else "Normal",
        ppr_xml=deepcopy(ppr) if ppr is not None else None,
        rpr_xml=deepcopy(rpr) if rpr is not None else None,
    )


def _clear_body_keep_sectpr(doc: Document) -> None:
    body = doc._element.body  # type: ignore[attr-defined]
    for child in list(body):
        # Keep section properties if present
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _add_para(doc: Document, proto: Proto, text: str, *, keep_whitespace: bool = False) -> None:
    p = doc.add_paragraph("")
    if proto.p_style:
        try:
            p.style = doc.styles[proto.p_style]
        except KeyError:
            log.warning("样式 '%s' 在输出文档中不存在，已降级为默认样式", proto.p_style)

    if proto.ppr_xml is not None:
        # Replace paragraph properties
        ppr = p._p.get_or_add_pPr()  # pylint: disable=protected-access
        p._p.remove(ppr)  # pylint: disable=protected-access
        p._p.insert(0, deepcopy(proto.ppr_xml))  # pylint: disable=protected-access

    run = p.add_run(text if keep_whitespace else text.strip())
    if proto.rpr_xml is not None:
        rpr = run._r.get_or_add_rPr()  # pylint: disable=protected-access
        run._r.remove(rpr)  # pylint: disable=protected-access
        run._r.insert(0, deepcopy(proto.rpr_xml))  # pylint: disable=protected-access


def _format_cn_date(d: date) -> str:
    return f"{d.year}年{d.month}月{d.day}日"


def main() -> None:
    ap = argparse.ArgumentParser(description="Generate docx from a .dotx fixed-format template and JSON content.")
    ap.add_argument("--template", required=True, type=Path, help="Path to .dotx template")
    ap.add_argument("--content", required=True, type=Path, help="Path to content JSON")
    ap.add_argument("--out", required=True, type=Path, help="Output .docx path")
    args = ap.parse_args()

    _validate_inputs(args.template, args.content)
    log.info("模板：%s", args.template)
    log.info("内容：%s", args.content)

    with tempfile.TemporaryDirectory(prefix="atc_dotx_") as _tmpdir:
        template_docx = _dotx_to_docx(args.template, Path(_tmpdir))
        tpl = Document(str(template_docx))

    # Pick paragraph prototypes by matching typical patterns from the provided template:
    # title: first non-empty paragraph
    title_p = next(_iter_nonempty_paras(tpl))
    title_idx = _find_index(tpl, title_p)

    to_p = _find_first(tpl, lambda p: p.text.strip().endswith("："))
    sec_h1_p = _find_first(tpl, lambda p: "、" in p.text and p.text.strip()[0] in "一二三四五六七八九十")
    subitem_p = _find_first(tpl, lambda p: p.text.strip().startswith("（") and "）" in p.text)
    closing_p = _find_first(tpl, lambda p: "妥否" in p.text or "请批示" in p.text)
    attach_head_p = _find_first(tpl, lambda p: p.text.strip().startswith("附件"))
    attach_head_idx = _find_index(tpl, attach_head_p)
    attach_item_p = _find_first(
        tpl,
        lambda p: _find_index(tpl, p) > attach_head_idx and re.match(r"^\d+\.", p.text.strip()) is not None,
    )
    sig_org_p = _find_first(tpl, lambda p: p.alignment == WD_ALIGN_PARAGRAPH.RIGHT)
    sig_org_idx = _find_index(tpl, sig_org_p)
    sig_date_p = _find_first(
        tpl,
        lambda p: _find_index(tpl, p) > sig_org_idx
        and ("年" in p.text and "月" in p.text and "日" in p.text),
    )

    proto_title = _proto_from_para(title_p)
    proto_to = _proto_from_para(to_p)
    proto_h1 = _proto_from_para(sec_h1_p)
    proto_sub = _proto_from_para(subitem_p)
    proto_body = _proto_from_para(
        _find_first(
            tpl,
            # Don't rely on wrapper identity; compare by paragraph index.
            lambda p: _find_index(tpl, p) != title_idx
            and p.style is not None
            and p.style.name == "Normal"
            and p.alignment is None
            and p.paragraph_format.first_line_indent is not None,
        )
    )
    proto_close = _proto_from_para(closing_p)
    proto_attach_head = _proto_from_para(attach_head_p)
    proto_attach_item = _proto_from_para(attach_item_p)
    proto_sig_org = _proto_from_para(sig_org_p)
    proto_sig_date = _proto_from_para(sig_date_p)

    content = _load_json(args.content)

    # Start from template to keep page setup/styles, but replace body content.
    _clear_body_keep_sectpr(tpl)

    _add_para(tpl, proto_title, str(content.get("title", "")), keep_whitespace=True)
    _add_para(tpl, proto_body, "", keep_whitespace=True)  # blank line

    to_line = str(content.get("to", "")).strip()
    if to_line:
        _add_para(tpl, proto_to, to_line, keep_whitespace=True)

    for s in content.get("preamble", []) or []:
        _add_para(tpl, proto_body, str(s), keep_whitespace=True)

    for sec in content.get("sections", []) or []:
        if sec.get("h1"):
            _add_para(tpl, proto_h1, str(sec["h1"]), keep_whitespace=True)
        if sec.get("h2"):
            _add_para(tpl, proto_sub, str(sec["h2"]), keep_whitespace=True)
        if sec.get("h3"):
            _add_para(tpl, proto_sub, str(sec["h3"]), keep_whitespace=True)
        for s in sec.get("paras", []) or []:
            _add_para(tpl, proto_body, str(s), keep_whitespace=True)

    closing = str(content.get("closing", "")).strip()
    if closing:
        _add_para(tpl, proto_close, closing, keep_whitespace=True)

    atts = content.get("attachments", []) or []
    if atts:
        _add_para(tpl, proto_body, "", keep_whitespace=True)
        _add_para(tpl, proto_attach_head, "附件：1." + str(atts[0]), keep_whitespace=True)
        for idx, item in enumerate(atts[1:], start=2):
            _add_para(tpl, proto_attach_item, f"{idx}." + str(item), keep_whitespace=True)

    # Match the template: leave a few blank lines before signature
    _add_para(tpl, proto_body, "", keep_whitespace=True)
    _add_para(tpl, proto_body, "", keep_whitespace=True)
    _add_para(tpl, proto_body, "", keep_whitespace=True)

    sig = content.get("signature", {}) or {}
    org = str(sig.get("org", "")).strip()
    dt_raw = str(sig.get("date", "")).strip()

    if org:
        # Match the template's visual positioning: the signature org line uses padding spaces
        # in addition to right alignment, so the org sits directly above the date line.
        _add_para(tpl, proto_sig_org, (" " * 5) + org + (" " * 7), keep_whitespace=True)

    if dt_raw:
        # accept ISO yyyy-mm-dd
        try:
            parts = dt_raw.split("-")
            if len(parts) != 3:
                raise ValueError(f"日期格式错误，应为 yyyy-mm-dd，当前：{dt_raw!r}")
            y, m0, d0 = [int(x) for x in parts]
            dt = date(y, m0, d0)
            _add_para(tpl, proto_sig_date, _format_cn_date(dt), keep_whitespace=True)
        except ValueError as e:
            log.warning("日期解析失败（%s），将直接使用原始字符串", e)
            _add_para(tpl, proto_sig_date, dt_raw, keep_whitespace=True)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(str(args.out))


if __name__ == "__main__":
    main()
